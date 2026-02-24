import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

# --- Data ---
NAME = "Alan Nascimento Batista Cordeiro"
TITLE = "Engenheiro de Dados | Analista de Business Intelligence | Automação de TI"
CONTACT = [
    "Santo André, SP",
    "+55 (11) 97055-6565",
    "eualannascimento@gmail.com",
    "linkedin.com/in/eualannascimento"
]

SUMMARY = (
    "Com 11 anos de trajetória no Bradesco, construí uma carreira sólida que evoluiu da linha de frente do atendimento ao cliente "
    "para a Engenharia de Dados. Minha jornada reflete um crescimento pautado pela resolução de problemas reais através da tecnologia, "
    "migrando da infraestrutura de TI para a Business Intelligence e análise avançada. Foco na criação de ativos digitais "
    "escaláveis, utilizando Databricks, Python e SQL para gerar valor real e integridade para o negócio."
)

EXPERIENCES = [
    {
        "company": "Bradesco",
        "dept": "Gestão Corp. de Riscos - GCR",
        "title": "Engenheiro de Dados",
        "period": "05/2025 - Atual",
        "bullets": [
            "Arquitetura Medalhão (Bronze/Silver/Gold) em Databricks (PySpark) e SQL para fluxos de Riscos.",
            "Processamento de grandes volumes de dados de fontes diversas (Hive, Teradata, Oracle, Azure Data Lake).",
            "Otimização de rotinas de ELT garantindo atualização diária (D-1) e alta disponibilidade dos dados.",
            "Governança técnica: documentação de metadados, linhagem de dados e controle de versão (Git)."
        ]
    },
    {
        "company": "Bradesco",
        "dept": "Vendas Digitais - VD",
        "title": "Analista de Business Intelligence",
        "period": "03/2021 - 04/2025",
        "bullets": [
            "Desenvolvimento de +15 dashboards corporativos (Power BI, Tableau, Looker) com modelagem Star Schema.",
            "Desenvolvimento de soluções low-code (PowerApps/Power Automate) para gestão de demandas internas.",
            "Liderança técnica e acompanhamento de equipe de dados, garantindo padrões de qualidade e documentação."
        ]
    },
    {
        "company": "Bradesco",
        "dept": "Infraestrutura de TI - DITI",
        "title": "Analista de Suporte de TI",
        "period": "02/2019 - 03/2021",
        "bullets": [
            "Suporte técnico local e remoto a parque de 3000+ estações utilizando escala e automação via PowerShell.",
            "Desenvolvimento de bibliotecas CLI para orquestração de scripts de manutenção e inventário.",
            "Redução estratégica de chamados através de análise de causa raiz e implementação de correções em massa."
        ]
    },
    {
        "company": "Bradesco",
        "dept": "Canais Digitais - DCD",
        "title": "Analista de Projetos",
        "period": "03/2016 - 02/2019",
        "bullets": [
            "Otimização de fluxos operacionais do SAC e Ouvidoria, elevando taxas de resolução e eficiência.",
            "Gestão de stakeholders e produção de reports analíticos periódicos para a Alta Liderança."
        ]
    },
    {
        "company": "Bradesco",
        "dept": "Canais Digitais - DCD",
        "title": "Analista de Atendimento",
        "period": "08/2014 - 03/2016",
        "bullets": [
            "Suporte transacional especializado focado em resolução de problemas e conformidade regulatória."
        ]
    }
]

SKILLS = [
    {"category": "Engenharia de Dados", "tools": "SQL (HQL/Teradata/Oracle), Python (Pandas/PySpark), Databricks, Hive, Git."},
    {"category": "Business Intelligence", "tools": "Power BI (DAX/M), Tableau, Looker Studio, Google Analytics, Excel Avançado/VBA."},
    {"category": "Automação & Sistemas", "tools": "PowerApps, Power Automate, PowerShell, Bash, Linux, Javascript, VBScript."},
    {"category": "Gestão & Ferramentas", "tools": "Scrum, COBIT, ITIL, SharePoint, Confluence, Notion, MS Office, Figma."}
]

EDUCATION = [
    "Pós-graduação em Engenharia de Dados - Data Science Academy (Conclusão Prevista 2026)",
    "Tecnólogo em Análise e Desenvolvimento de Sistemas - FATEC Ipiranga (Graduado 2018)",
    "Técnico em Desenvolvimento Web - ETEC Camargo Aranha (Graduado 2012)"
]

CERTIFICATIONS = [
    "Azure Fundamentals (AZ-900) - Microsoft (2026)",
    "Databricks Fundamentals - Databricks (2025)",
    "Notion Workflows - Notion (2025)",
    "Notion Essentials - Notion (2025)",
    "Scrum Foundation Professional (SFPC) (2020)",
    "COBIT 5 Foundation - APMG International (2019)"
]

# --- DOCX Generation ---
def generate_docx(filename):
    doc = Document()
    s = doc.sections[0]
    # Comfortable margins for 1-page fit
    s.top_margin, s.bottom_margin = Inches(0.2), Inches(0.2)
    s.left_margin, s.right_margin = Inches(0.3), Inches(0.3)

    # Header
    h_para = doc.add_paragraph()
    h_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h_para.add_run(NAME.upper())
    run.bold, run.font.size = True, Pt(16) # Increased
    
    t_para = doc.add_paragraph()
    t_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t_para.add_run(TITLE)
    run.font.size = Pt(10) # Increased
    run.italic = True

    c_para = doc.add_paragraph()
    c_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = c_para.add_run(" | ".join(CONTACT))
    run.font.size = Pt(9) # Increased
    doc.add_paragraph().paragraph_format.space_after = Pt(4) # Increased

    def section_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold, run.font.size = True, Pt(11) # Increased
        p.paragraph_format.space_before = Pt(2) # Increased
        p.paragraph_format.space_after = Pt(2) # Increased

    section_heading("RESUMO PROFISSIONAL")
    p = doc.add_paragraph(SUMMARY)
    p.paragraph_format.line_spacing = 1 # Comfortable
    p.paragraph_format.space_after = Pt(3) # Increased

    section_heading("EXPERIÊNCIA PROFISSIONAL")
    for exp in EXPERIENCES:
        # Title (Bold)
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(exp['title'])
        run_title.bold, run_title.font.size = True, Pt(9) # Increased
        p_title.paragraph_format.space_before = Pt(3) # Increased
        p_title.paragraph_format.space_after = Pt(0)
        
        # Company - Dept (Next line) + Period on Right
        p_info = doc.add_paragraph()
        p_info.paragraph_format.space_after = Pt(2) # Increased
        run_info = p_info.add_run(f"{exp['company']} - {exp['dept']}")
        run_info.italic, run_info.font.size = True, Pt(9.5) # Increased
        p_info.add_run(f" \t {exp['period']}").italic = True
        
        for bullet in exp['bullets']:
            bp = doc.add_paragraph(bullet, style='List Bullet')
            bp.paragraph_format.line_spacing = 1 # Increased
            bp.paragraph_format.space_after = Pt(1) # Increased
            bp.paragraph_format.left_indent = Inches(0.25)
        
    section_heading("COMPETÊNCIAS TÉCNICAS")
    for skill in SKILLS:
        p = doc.add_paragraph(f"{skill['category']}: {skill['tools']}", style='List Bullet')
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.25) # Match indentation of other bullets
        for run in p.runs:
            run.font.size = Pt(9.5)

    section_heading("FORMAÇÃO ACADÊMICA")
    for edu in EDUCATION:
        p = doc.add_paragraph(edu, style='List Bullet')
        p.paragraph_format.space_after = Pt(1) # Increased
        for run in p.runs: run.font.size = Pt(9.5) # Increased

    section_heading("CERTIFICAÇÕES")
    for cert in CERTIFICATIONS:
        p = doc.add_paragraph(cert, style='List Bullet')
        p.paragraph_format.space_after = Pt(1) # Increased
        for run in p.runs:
            run.bold = False
            run.font.size = Pt(9.5) # Increased

    doc.save(filename)

# --- PDF Generation ---
class PDF(FPDF):
    def header(self):
        self.set_font('helvetica', 'B', 16)
        self.cell(0, 8, NAME.upper(), ln=True, align='C')
        self.set_font('helvetica', 'I', 10)
        self.cell(0, 5, TITLE, ln=True, align='C')
        self.set_font('helvetica', '', 9.5)
        self.cell(0, 5, " | ".join(CONTACT), ln=True, align='C')
        self.ln(2)

    def section_title(self, label):
        self.ln(2)
        self.set_font('helvetica', 'B', 11)
        self.set_fill_color(242, 242, 242)
        self.cell(0, 7, "  " + label.upper(), ln=True, fill=True)
        self.ln(1.5) # Reduced from 2

def generate_pdf(filename):
    pdf = PDF()
    pdf.add_page()
    # Comfortable margins
    pdf.set_margins(12, 10, 12)
    
    # Summary
    pdf.section_title("Resumo Profissional")
    pdf.set_font('helvetica', '', 9.5)
    pdf.multi_cell(0, 4.5, SUMMARY)
    
    # Experience
    pdf.section_title("Experiência Profissional")
    for exp in EXPERIENCES:
        # Title
        pdf.set_font('helvetica', 'B', 10.5)
        pdf.cell(0, 5, exp['title'], ln=True)
        
        # Company - Dept + Period
        pdf.set_font('helvetica', 'I', 9.5)
        pdf.cell(155, 5, f"{exp['company']} - {exp['dept']}")
        pdf.cell(0, 5, exp['period'], ln=True, align='R')
        
        # Bullets
        pdf.set_font('helvetica', '', 9.5)
        for bullet in exp['bullets']:
            pdf.set_x(18)
            pdf.cell(4, 4.5, chr(149), ln=0)
            pdf.multi_cell(0, 4.2, bullet) # Slight reduction in line height (4.5 -> 4.2)
        pdf.ln(1)

    # Skills
    pdf.section_title("Competências Técnicas")
    pdf.set_font('helvetica', '', 9.5) # No bold, normal font
    for skill in SKILLS:
        pdf.set_x(18)
        pdf.cell(4, 4.5, chr(149), ln=0)
        pdf.multi_cell(0, 4.5, f"{skill['category']}: {skill['tools']}")
    
    # Education
    pdf.section_title("Formação Acadêmica")
    pdf.set_font('helvetica', '', 9.5)
    for edu in EDUCATION:
        pdf.set_x(18)
        pdf.cell(4, 5, chr(149), ln=0)
        pdf.cell(0, 5, edu, ln=True)
    
    # Certifications
    pdf.section_title("Certificações")
    pdf.set_font('helvetica', '', 9.5)
    for cert in CERTIFICATIONS:
        pdf.set_x(18)
        pdf.cell(4, 5, chr(149), ln=0)
        pdf.cell(0, 5, cert, ln=True)

    pdf.output(filename)

if __name__ == "__main__":
    generate_docx("curriculo_alan.docx")
    generate_pdf("curriculo_alan.pdf")
    generate_pdf("curriculo_alan2.pdf")
